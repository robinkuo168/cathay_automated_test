# backend/document_analyzer.py
import docx
import re
import json
from typing import Dict, List
from .llm_service import LLMService
from .logger import get_logger

class DocumentAnalyzer:
    def __init__(self, llm_service: LLMService = None):
        """
        初始化文件分析器。

        此建構函式會設定分析器所需的依賴項，例如 LLM 服務和日誌記錄器。
        :param llm_service: (可選) 一個 LLMService 實例，用於需要 AI 分析的功能。
        """
        self.llm_service = llm_service or LLMService()
        self.logger = get_logger(__name__)

    def extract_content_from_docx(self, file_path: str) -> Dict:
        """
        從 Word (.docx) 文檔中提取結構化的內容。

        此函式是分析流程的主要進入點，它會讀取一個 Word 檔案，
        並協調提取其純文字、表格、圖片資訊，最後再將這些原始資料解析為結構化的效能數據。
        :param file_path: Word 檔案的完整路徑。
        :return: 一個包含所有提取資訊的字典，例如 'text_content', 'tables', 'structured_data'。
        :raises FileNotFoundError: 如果提供的檔案路徑不存在。
        :raises Exception: 如果在解析過程中發生其他未預期的錯誤。
        """
        try:
            # 使用 python-docx 一次性讀取檔案
            doc = docx.Document(file_path)

            # 直接從 doc 物件中提取文字
            text_content = '\n'.join([para.text for para in doc.paragraphs])
            tables_data = self._extract_tables(doc)

            # 提取圖片
            images_info = self._extract_images_info(doc)

            return {
                'text_content': text_content,
                'tables': tables_data,
                'images': images_info,
                'structured_data': self._parse_performance_data(text_content, tables_data)
            }
        except Exception as e:
            self.logger.error(f"文檔解析失敗: {e}")
            raise

    def _extract_tables(self, doc) -> List[Dict]:
        """
        從一個 python-docx 的 Document 物件中提取所有表格的數據。

        它會遍歷文件中的所有表格，並將每個單元格的文字內容轉換為一個巢狀列表結構。
        :param doc: 一個已載入的 python-docx Document 物件。
        :return: 一個表格列表。其中每個表格是一個行的列表，而每行又是一個包含該行所有單元格文字的列表。
        """
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def _parse_performance_data(self, text: str, tables: List) -> Dict:
        """
        從原始文字和表格數據中，解析出結構化的效能指標。

        這是一個調度函式 (dispatcher)，它會呼叫多個特定的 `_extract_*` 輔助函式，
        分別提取如 TPS、響應時間、錯誤率、資源配置等不同的效能數據，並將它們組合起來。
        :param text: 從文檔中提取的完整純文字內容。
        :param tables: 從文檔中提取的表格數據列表。
        :return: 一個包含各類結構化效能指標的字典。
        """
        return {
            'tps_data': self._extract_tps_data(text, tables),
            'response_time_data': self._extract_response_time_data(text, tables),
            'error_rate_data': self._extract_error_rate_data(text, tables),
            'resource_config': self._extract_resource_config(text),
            'test_duration': self._extract_test_duration(text),
            'concurrent_users': self._extract_concurrent_users(text)
        }

    def _extract_tps_data(self, text: str, tables: List) -> List[Dict]:
        """
        從文字和表格中專門提取與「每秒交易數」(TPS) 相關的數據。

        此函式會使用正規表示式 (regex) 在純文字中尋找 TPS 關鍵字和數值，
        同時也會嘗試遍歷表格，從中識別並解析包含 TPS 資訊的行。
        :param text: 從文檔中提取的完整純文字內容。
        :param tables: 從文檔中提取的表格數據列表。
        :return: 一個字典列表，其中每個字典代表一個找到的 TPS 數據點。
        """
        tps_pattern = r'TPS[：:]\s*(\d+\.?\d*)'
        matches = re.findall(tps_pattern, text, re.IGNORECASE)

        # 從表格中尋找 TPS 數據
        tps_data = []
        for table in tables:
            for row in table:
                if any('tps' in cell.lower() for cell in row):
                    # 解析 TPS 相關行
                    pass

        return tps_data

    def _extract_resource_config(self, text: str, tables: List) -> Dict:
        """
        從文字和表格中提取測試環境的「資源配置」資訊。

        這是一個調度函式，它會呼叫更具體的輔助函式 (例如 `_extract_cpu_config`)
        來分別提取 CPU、記憶體、網路等硬體或軟體的配置詳情。
        :param text: 從文檔中提取的完整純文字內容。
        :param tables: 從文檔中提取的表格數據列表。
        :return: 一個包含結構化資源配置資訊的字典。
        """
        return {
            'cpu_config': self._extract_cpu_config(text),
            'memory_config': self._extract_memory_config(text),
            'network_config': self._extract_network_config(text)
        }
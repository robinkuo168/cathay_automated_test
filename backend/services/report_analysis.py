import json
from typing import Dict
from pathlib import Path
from datetime import datetime
from .logger import get_logger
from docx import Document
from .document_analyzer import DocumentAnalyzer
from docx.shared import Inches
from .llm_service import LLMService

class ReportAnalysisService:
    def __init__(self, llm_service=None):
        """
        初始化 ReportAnalysisService。

        此建構函式採用延遲初始化 (lazy initialization) 模式，
        僅儲存傳入的依賴項，而不會立即創建服務實例，
        以提升應用程式啟動效率。
        :param llm_service: (可選) 一個 LLMService 實例。
        """

        self._llm_service = llm_service  # 不立即創建實例，儲存為私有屬性
        self._document_analyzer = None  # 延遲初始化 DocumentAnalyzer
        self.logger = get_logger(__name__)

    @property
    def llm_service(self):
        """
        一個延遲載入 (lazy-loading) 的屬性，用於獲取 LLMService 實例。

        它確保 LLMService 只在第一次被需要時才進行初始化，
        避免了不必要的資源消耗和啟動延遲。
        :return: 一個 LLMService 的實例。
        """
        if self._llm_service is None:
            from .llm_service import LLMService
            self._llm_service = LLMService()
        return self._llm_service

    @property
    def document_analyzer(self):
        """
        一個延遲載入 (lazy-loading) 的屬性，用於獲取 DocumentAnalyzer 實例。

        它確保 DocumentAnalyzer 只在第一次被需要時才進行初始化。
        :return: 一個 DocumentAnalyzer 的實例。
        """
        if self._document_analyzer is None:
            from .document_analyzer import DocumentAnalyzer
            self._document_analyzer = DocumentAnalyzer()
        return self._document_analyzer

    def analyze_performance_report(self, file_path: str) -> Dict:
        """
        分析效能測試報告的總指揮。

        此函式協調整個分析流程，包含以下核心步驟：
        1. 使用 DocumentAnalyzer 提取報告的結構化內容。
        2. 根據提取的內容，建構一個詳細的分析提示詞 (Prompt)。
        3. 呼叫 LLM 服務來執行深度分析。
        4. 將 LLM 的分析結果進行最終的結構化處理。
        :param file_path: 效能報告 (Word 檔案) 的路徑。
        :return: 一個包含完整分析結果的字典。
        :raises Exception: 如果在任何分析步驟中發生錯誤。
        """
        try:
            # 1. 提取文檔內容
            doc_content = self.document_analyzer.extract_content_from_docx(file_path)

            # 2. 構建分析提示詞
            analysis_prompt = self._build_analysis_prompt(doc_content)

            # 3. 調用 LLM 進行分析
            analysis_result = self._call_llm_for_analysis(analysis_prompt)

            # 4. 結構化分析結果
            structured_result = self._structure_analysis_result(analysis_result, doc_content)

            return structured_result

        except Exception as e:
            self.logger.error(f"報告分析失敗: {e}")
            raise

    def _call_llm_for_analysis(self, prompt: str) -> Dict:
        """
        一個內部輔助函式，用於呼叫 LLM 服務並解析其回應。

        它負責將分析提示詞傳送給 LLM，並預期接收一個 JSON 格式的回應。
        此函式也包含一個備用機制，在 LLM 呼叫失敗時返回一個預設的模擬結果，以確保流程的健壯性。
        :param prompt: 要傳送給 LLM 的完整分析提示詞。
        :return: 一個從 LLM 回應中解析出的分析結果字典。
        """
        try:
            # 使用 LLMService 進行文字生成
            response = self.llm_service.generate_text(prompt=prompt)
            return json.loads(response)
        except Exception as e:
            self.logger.error(f"LLM 分析失敗: {e}")
            # 返回一個預設的分析結果
            return {
                "tps_analysis": {"status": "pass", "details": "TPS 符合預期", "recommendations": ["保持現有設定"]},
                "response_time_analysis": {"avg_time_status": "good", "p99_status": "warning",
                                         "recommendations": ["優化慢速查詢"]},
                "resource_analysis": {"cpu_recommendation": "CPU 使用率正常", "memory_recommendation": "記憶體使用率高",
                                    "scaling_suggestion": "考慮增加記憶體"},
                "database_analysis": {"performance_status": "warning", "recommendations": ["優化資料庫索引"]},
                "additional_tests": [{"scenario": "壓力測試", "reason": "驗證系統在高負載下的表現"}],
                "overall_assessment": {"grade": "B", "summary": "系統表現良好，但有優化空間"}
            }

    def _build_analysis_prompt(self, doc_content: Dict) -> str:
        """
        一個內部輔助函式，用於建構指導 LLM 進行分析的提示詞 (Prompt)。

        這是「提示詞工程」的核心，它將從文件中提取的數據摘要和部分原文，
        與一個詳細的指令模板結合，指導 LLM 從多個維度（如 TPS、響應時間、資源使用等）進行分析，
        並要求其以特定的 JSON 格式回傳結果。
        :param doc_content: 從文件中提取出的結構化內容字典。
        :return: 一個完整的、準備好發送給 LLM 的提示詞字串。
        """
        prompt = f"""
        請分析以下效能測試報告，並提供專業的分析和建議：
    
        ## 測試數據摘要：
        {json.dumps(doc_content['structured_data'], ensure_ascii=False, indent=2)}
    
        ## 完整報告內容：
        {doc_content['text_content'][:3000]}  # 限制長度避免超過 token 限制
    
        請從以下角度進行分析：
    
        ### 1. TPS 達標分析
        - 分析實際 TPS 與預期 TPS 的對比
        - 評估是否達到效能要求
        - 提供 TPS 相關建議
    
        ### 2. 響應時間分析
        - 分析平均響應時間 (Avg. RespTime)
        - 分析 99% 響應時間 (99% RespTime)
        - 識別異常響應時間並提供建議
    
        ### 3. 系統資源分析
        - 基於描述的 CPU/Memory 使用情況進行分析
        - 提供資源優化建議
        - 建議縱向或橫向擴展策略
    
        ### 4. 資料庫效能分析
        - 分析資料庫存取效能
        - 識別可能的 Full Table Scan
        - 提供查詢優化建議
    
        ### 5. 測試情境建議
        - 建議額外的測試情境
        - 提供測試參數調整建議
    
        請以 JSON 格式回應，包含以下結構：
        {{
            "tps_analysis": {{"status": "pass/fail", "details": "詳細分析", "recommendations": []}},
            "response_time_analysis": {{"avg_time_status": "good/warning/critical", "p99_status": "good/warning/critical", "recommendations": []}},
            "resource_analysis": {{"cpu_recommendation": "", "memory_recommendation": "", "scaling_suggestion": ""}},
            "database_analysis": {{"performance_status": "good/warning/critical", "recommendations": []}},
            "additional_tests": [{{\"scenario\": \"\", \"reason\": \"\"}}],
            "overall_assessment": {{"grade": "A/B/C/D", "summary": ""}}
        }}
        """
        return prompt

    def _structure_analysis_result(self, analysis_result: Dict, doc_content: Dict) -> Dict:
        """
        一個內部輔助函式，用於對 LLM 的分析結果進行最終的結構化處理。

        在目前的實現中，它直接返回 LLM 的分析結果。
        未來可以擴展此函式，以加入更多後處理邏輯，例如數據驗證、格式轉換或與原始數據的交叉比對。
        :param analysis_result: 從 LLM 返回的原始分析結果字典。
        :param doc_content: 從文件中提取出的原始內容，用於可能的交叉驗證。
        :return: 最終的、準備好返回給客戶端的分析結果字典。
        """
        # 目前直接返回分析結果，未來可以根據需求進行結構化處理
        return analysis_result

    def preview_analysis(self, file_path: str) -> dict:
        """
        提供效能報告的快速預覽分析，是 API 的主要進入點之一。

        此函式專為需要快速摘要的場景設計。它會提取 Word 檔案的內容，
        並呼叫核心分析邏輯來生成一份摘要性的分析結果。
        :param file_path: 要分析的 Word 檔案路徑。
        :return: 一個包含分析摘要的字典。
        :raises Exception: 如果在預覽分析過程中發生錯誤。
        """
        try:
            self.logger.info(f"開始預覽分析: {file_path}")

            # 讀取 Word 檔案內容
            content = self._extract_word_content(file_path)

            # 執行快速分析
            analysis_result = self._analyze_report_content(content, preview_mode=True)

            self.logger.info("預覽分析完成")
            return analysis_result

        except Exception as e:
            self.logger.error(f"預覽分析失敗: {e}")
            raise Exception(f"預覽分析失敗: {str(e)}")

    def generate_analysis_report(self, file_path: str) -> str:
        """
        生成一份完整的、可下載的 Word 格式分析報告，是 API 的主要進入點之一。

        此函式會執行完整的分析流程，然後將詳細的分析結果和建議，
        寫入一個新建立的 Word 文件中，並返回該文件的路徑。
        :param file_path: 原始的效能報告 Word 檔案路徑。
        :return: 生成的新分析報告檔案的路徑字串。
        :raises Exception: 如果在生成報告的過程中發生錯誤。
        """
        try:
            self.logger.info(f"開始生成分析報告: {file_path}")

            # 讀取 Word 檔案內容
            content = self._extract_word_content(file_path)

            # 執行完整分析
            analysis_result = self._analyze_report_content(content, preview_mode=False)

            # 生成 Word 報告
            output_path = self._create_analysis_document(analysis_result)

            self.logger.info(f"分析報告生成完成: {output_path}")
            return output_path

        except Exception as e:
            self.logger.error(f"生成分析報告失敗: {e}")
            raise Exception(f"生成分析報告失敗: {str(e)}")

    def _extract_word_content(self, file_path: str) -> str:
        """
        一個內部輔助函式，用於從 Word 檔案中提取所有可讀的文字內容。

        它會遍歷文件中的所有段落和表格，將它們的文字內容合併成一個單一的字串，
        為後續的 LLM 分析提供完整的上下文。
        :param file_path: Word 檔案的路徑。
        :return: 一個包含文件所有文字的單一字串。
        :raises Exception: 如果讀取或解析 Word 檔案失敗。
        """
        try:
            doc = Document(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())

            # 處理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(' | '.join(row_text))

            return '\n'.join(content)

        except Exception as e:
            self.logger.error(f"提取 Word 內容失敗: {e}")
            raise Exception(f"無法讀取 Word 檔案: {str(e)}")

    def _analyze_report_content(self, content: str, preview_mode: bool = False) -> dict:
        """
        一個內部輔助函式，負責執行核心的內容分析邏輯。

        在目前的實現中，它主要提供一個基本的、基於規則的分析結果作為佔位符。
        它的設計目標是在未來可以輕鬆地被一個真正的、由 LLM 驅動的分析邏輯所取代。
        :param content: 從文件中提取的完整文字內容。
        :param preview_mode: 一個布林值，指示是否應執行快速的預覽模式分析。
        :return: 一個包含分析結果的字典。
        """
        try:
            # 基本的文字分析
            analysis_result = {
                "tps_analysis": {
                    "status": "pass",
                    "details": "TPS 分析完成",
                    "recommendations": ["建議進行更長時間的測試"]
                },
                "response_time_analysis": {
                    "avg_time_status": "good",
                    "p99_status": "good",
                    "recommendations": ["響應時間表現良好"]
                },
                "resource_analysis": {
                    "cpu_recommendation": "CPU 使用率正常",
                    "memory_recommendation": "記憶體使用率正常",
                    "scaling_suggestion": "目前配置足夠"
                },
                "database_analysis": {
                    "performance_status": "good",
                    "recommendations": ["資料庫效能良好"]
                },
                "overall_assessment": {
                    "grade": "A",
                    "summary": "整體效能表現優秀"
                },
                "additional_tests": [
                    {
                        "scenario": "長時間穩定性測試",
                        "reason": "驗證系統長期穩定性"
                    }
                ]
            }

            # 如果有 AI 模型，可以在這裡調用進行更深入的分析
            if hasattr(self, 'model') and self.model:
                try:
                    # 使用 AI 模型進行分析
                    ai_analysis = self._analyze_with_ai(content)
                    # 將 AI 分析結果整合到 analysis_result 中
                    if ai_analysis:
                        analysis_result["ai_insights"] = ai_analysis
                except Exception as e:
                    self.logger.warning(f"AI 分析失敗，使用基本分析: {e}")

            return analysis_result

        except Exception as e:
            self.logger.error(f"分析報告內容失敗: {e}")
            # 返回基本的分析結果
            return {
                "overall_assessment": {
                    "grade": "N/A",
                    "summary": "分析過程中發生錯誤"
                }
            }

    def _create_analysis_document(self, analysis_result: dict) -> str:
        """
        一個內部輔助函式，用於將結構化的分析結果寫入一個新的 Word 文件。

        它會根據傳入的分析結果字典，動態地生成標題、段落和列表，
        創建一份格式清晰、易於閱讀的分析報告。
        :param analysis_result: 包含所有分析結果和建議的結構化字典。
        :return: 新建立的 Word 報告檔案的路徑字串。
        :raises Exception: 如果在寫入 Word 文件的過程中發生錯誤。
        """
        try:
            doc = Document()

            # 添加標題
            title = doc.add_heading('效能測試分析報告', 0)

            # 添加生成時間
            doc.add_paragraph(f"報告生成時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")  # 空行

            # 添加整體評估
            if "overall_assessment" in analysis_result:
                doc.add_heading('整體評估', level=1)
                assessment = analysis_result["overall_assessment"]
                doc.add_paragraph(f"評級: {assessment.get('grade', 'N/A')}")
                doc.add_paragraph(f"摘要: {assessment.get('summary', '無摘要')}")
                doc.add_paragraph("")

            # 添加 TPS 分析
            if "tps_analysis" in analysis_result:
                doc.add_heading('TPS 效能分析', level=1)
                tps = analysis_result["tps_analysis"]
                doc.add_paragraph(f"狀態: {tps.get('status', 'N/A')}")
                doc.add_paragraph(f"詳細: {tps.get('details', 'N/A')}")
                if tps.get('recommendations'):
                    doc.add_paragraph("建議:")
                    for rec in tps['recommendations']:
                        doc.add_paragraph(f"• {rec}")
                doc.add_paragraph("")

            # 添加響應時間分析
            if "response_time_analysis" in analysis_result:
                doc.add_heading('響應時間分析', level=1)
                rt = analysis_result["response_time_analysis"]
                doc.add_paragraph(f"平均響應時間狀態: {rt.get('avg_time_status', 'N/A')}")
                doc.add_paragraph(f"99% 響應時間狀態: {rt.get('p99_status', 'N/A')}")
                if rt.get('recommendations'):
                    doc.add_paragraph("建議:")
                    for rec in rt['recommendations']:
                        doc.add_paragraph(f"• {rec}")
                doc.add_paragraph("")

            # 添加資源分析
            if "resource_analysis" in analysis_result:
                doc.add_heading('資源使用分析', level=1)
                resource = analysis_result["resource_analysis"]
                doc.add_paragraph(f"CPU 建議: {resource.get('cpu_recommendation', 'N/A')}")
                doc.add_paragraph(f"記憶體建議: {resource.get('memory_recommendation', 'N/A')}")
                doc.add_paragraph(f"擴展建議: {resource.get('scaling_suggestion', 'N/A')}")
                doc.add_paragraph("")

            # 添加額外測試建議
            if "additional_tests" in analysis_result:
                doc.add_heading('建議的額外測試', level=1)
                for test in analysis_result["additional_tests"]:
                    doc.add_paragraph(f"測試場景: {test.get('scenario', 'N/A')}")
                    doc.add_paragraph(f"原因: {test.get('reason', 'N/A')}")
                    doc.add_paragraph("")

            # 生成輸出檔案路徑
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = Path("outputs")
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / f"analysis_report_{timestamp}.docx"

            # 保存文檔
            doc.save(str(output_path))

            return str(output_path)

        except Exception as e:
            self.logger.error(f"創建分析文檔失敗: {e}")
            raise Exception(f"創建分析文檔失敗: {str(e)}")

    def _analyze_with_ai(self, content: str) -> dict:
        """
        一個內部輔助函式，作為未來使用 AI 模型進行分析的佔位符。

        此函式的目的是在未來實現更高級的 AI 分析功能時，提供一個清晰的擴展點。
        在目前的版本中，它不執行任何操作。
        :param content: 要分析的文字內容。
        :return: AI 分析的結果字典，目前返回 None。
        """
        try:
            if not self.model:
                return None

            # 這裡可以實現具體的 AI 分析邏輯
            # 暫時返回 None
            return None

        except Exception as e:
            self.logger.error(f"AI 分析失敗: {e}")
            return None
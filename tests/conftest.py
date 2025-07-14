"""
測試配置和共享夾件
"""
import pytest
from fastapi.testclient import TestClient
from backend.main import app
import os
import shutil
from pathlib import Path
from dotenv import load_dotenv
from unittest.mock import patch, MagicMock
import sys

# 在導入任何後端代碼之前，先模擬 ibm_watsonx_ai 套件
sys.modules['ibm_watsonx_ai'] = MagicMock()
sys.modules['ibm_watsonx_ai.foundation_models'] = MagicMock()
sys.modules['ibm_watsonx_ai.foundation_models.ModelInference'] = MagicMock()

# 加載測試環境變數
load_dotenv('.env.test')

# 測試用的臨時目錄
TEST_UPLOAD_DIR = "test_uploads"
TEST_OUTPUT_DIR = "test_outputs"

@pytest.fixture(scope="session", autouse=True)
def setup_test_environment():
    """設置測試環境"""
    # 創建測試用臨時目錄
    os.makedirs(TEST_UPLOAD_DIR, exist_ok=True)
    os.makedirs(TEST_OUTPUT_DIR, exist_ok=True)
    
    yield  # 測試執行
    
    # 測試結束後清理
    if os.path.exists(TEST_UPLOAD_DIR):
        shutil.rmtree(TEST_UPLOAD_DIR)
    if os.path.exists(TEST_OUTPUT_DIR):
        shutil.rmtree(TEST_OUTPUT_DIR)

@pytest.fixture
def test_app():
    """創建測試客戶端"""
    with TestClient(app) as test_client:
        # 設置測試模式
        with patch('backend.main.app') as mock_app:
            mock_app.config = {'TESTING': True}
            yield test_client

@pytest.fixture
def sample_docx_file():
    """創建一個測試用的 DOCX 文件"""
    from docx import Document
    import io
    
    doc = Document()
    doc.add_paragraph("這是一個測試文檔")
    doc.add_paragraph("包含一些測試內容")
    
    file_path = os.path.join(TEST_UPLOAD_DIR, "test_document.docx")
    doc.save(file_path)
    
    # 返回文件路徑和二進制內容
    with open(file_path, "rb") as f:
        content = f.read()
    
    return file_path, content

@pytest.fixture
def sample_jmx_content():
    """返回一個基本的 JMX 文件內容"""
    return """<?xml version="1.0" encoding="UTF-8"?>
<jmeterTestPlan version="1.2" properties="5.0" jmeter="5.4.1">
  <hashTree>
    <TestPlan guiclass="TestPlanGui" testclass="TestPlan" testname="Test Plan" enabled="true">
      <boolProp name="TestPlan.functional_mode">false</boolProp>
      <elementProp name="TestPlan.user_defined_variables" elementType="Arguments" guiclass="ArgumentsPanel" testclass="Arguments" testname="User Defined Variables" enabled="true">
        <collectionProp name="Arguments.arguments"/>
      </elementProp>
    </TestPlan>
    <hashTree/>
  </hashTree>
</jmeterTestPlan>
"""

@pytest.fixture
def mock_llm_service():
    """模擬 LLM 服務"""
    mock_service = MagicMock()
    mock_service.generate_response.return_value = """```xml
    <?xml version="1.0" encoding="UTF-8"?>
    <jmeterTestPlan>
        <hashTree>
            <TestPlan testname="Test Plan"/>
        </hashTree>
    </jmeterTestPlan>
    ```"""
    mock_service.generate_structured_response.return_value = {
        "summary": "測試摘要",
        "metrics": {"throughput": 100, "error_rate": 0.1},
        "recommendations": ["建議1", "建議2"]
    }
    
    # 模擬所有 LLM 服務相關的導入和初始化
    with patch('backend.services.llm_service.LLMService', return_value=mock_service):
        with patch('backend.main.get_llm_service', return_value=mock_service):
            with patch.dict('sys.modules', {
                'ibm_watsonx_ai': MagicMock(),
                'ibm_watsonx_ai.foundation_models': MagicMock(),
                'ibm_watsonx_ai.foundation_models.ModelInference': MagicMock()
            }):
                yield mock_service
